import discord
from discord.ext import commands
from docx import Document

client = commands.Bot(command_prefix='$')


async def msg_compile(ctx, mlimit=None, document=None, chnl=None):
    if chnl is None:
        if mlimit is None:
            all_messages = await ctx.channel.history(oldest_first=True).flatten()
        else:
            input_mlimit = int(mlimit) + 1
            all_messages = await ctx.channel.history(limit=input_mlimit).flatten()
            all_messages.reverse()
        all_messages = all_messages[:-1]
        for message in all_messages:
            if document is None:
                await ctx.send(message.content)
            else:
                document.add_paragraph(message.content)
    else:
        if mlimit is None:
            all_messages = await chnl.history(oldest_first=True).flatten()
        else:
            input_mlimit = int(mlimit) + 1
            all_messages = await chnl.history(limit=input_mlimit).flatten()
            all_messages.reverse()
        for message in all_messages:
            if document is None:
                await ctx.send(message.content)
            else:
                document.add_paragraph(message.content)


async def log_msgs(ctx, mlimit=None, chnl=None):
    document = Document()
    document.add_paragraph('[Beginning of Log]')
    if chnl is None:
        await msg_compile(ctx, mlimit, document)
    else:
        await msg_compile(ctx, mlimit, document, chnl)
    document.add_paragraph('[End of Log]')
    document.save('D:/Coding_Projects/Discord_Logger_Bot/Temp/test.docx')
    temp_file = discord.File('D:/Coding_Projects/Discord_Logger_Bot/Temp/test.docx', filename='log.docx')
    await ctx.send('Here is your log!', file=temp_file)


async def clear_msgs(ctx, mlimit=None, chnl=None):
    user = ctx.author
    if chnl is None:
        if user.guild_permissions.administrator or user.guild_permissions.manage_messages:
            if mlimit is None:
                deleting = await ctx.channel.purge()
            else:
                deleting = await ctx.channel.purge(limit=int(mlimit) + 1)
            await ctx.send("Deleted " + str(len(deleting) - 1) + " messages from " + ctx.channel.name + ".")
        else:
            await ctx.send("no")
    else:
        if user.guild_permissions.administrator or user.guild_permissions.manage_messages:
            if mlimit is None:
                deleting = await chnl.purge()
            else:
                deleting = await chnl.purge(limit=int(mlimit))
            await ctx.send("Deleted " + str(len(deleting)) + " messages from " + chnl.name + ".")
        else:
            await ctx.send("no")


async def check_channel(ctx, chnl):
    for txt_channel in ctx.guild.text_channels:
        if chnl == txt_channel.name:
            loc = txt_channel;
    if loc is None:
        return None
    else:
        return loc


@client.event
async def on_ready():
    print('Bot is ready to roll!')
    game = discord.Game("with mere mortals")
    await client.change_presence(status=discord.Status.idle, activity=game)


@client.command()
async def copy(ctx, arg):
    await ctx.send(arg)


@client.command()
async def ping(ctx):
    await ctx.send(f"Pong! {round(client.latency * 1000)}ms")


@client.command()
async def username(ctx):
    try:
        await ctx.send(ctx.author.nick)
    except:
        await ctx.send(ctx.author.name)


@client.command()
async def channel(ctx, chnl=None):
    output = str(chnl)
    if chnl is None:
        await ctx.send(ctx.channel)
    elif output not in ctx.guild.text_channels:
        await ctx.send("That is not a channel in this server.")
    else:
        await ctx.send(output)


@client.command()
async def compile_messages(ctx, mlimit=None):
    await msg_compile(ctx, mlimit)


@client.command()
async def log(ctx, mlimit=None):
    await log_msgs(ctx, mlimit)


@client.command()
async def log_channel(ctx, chnl=None, mlimit=None):
    if chnl is None:
        await ctx.send("You must send a channel.")
    else:
        loc = await check_channel(ctx, chnl)
        if loc is None:
            await ctx.send("That channel does not exist.")
        else:
            await log_msgs(ctx, mlimit, loc)


@client.command()
async def clear(ctx, mlimit=None):
    await clear_msgs(ctx, mlimit)


@client.command()
async def clear_channel(ctx, chnl=None, mlimit=None):
    if chnl is None:
        await ctx.send("You must send a channel.")
    else:
        loc = await check_channel(ctx, chnl)
        if loc is None:
            await ctx.send("That channel does not exist.")
        else:
            await clear_msgs(ctx, mlimit, loc)


@client.command(aliases=['commands', 'helpme'])
async def _help(ctx):
    await ctx.send("```"
                   "$helpme or $commands - displays this command.\n\n"
                   "$copy ''x'' - the bot mimics whatever you type in quotes (or a single word without quotes).\n\n"
                   "$ping - the bot tells you your latency to the server.\n\n"
                   "$channel x - the bot tells you the name of the current channel, or if x channel exists in the "
                   "current server.\n\n"
                   "$username - sends your current nickname in this server, or if you don't have one, your username."
                   "\n\n"
                   "$log x - log the most recent x messages, or if you don't give it an x, log the entire channel.\n\n"
                   "$log_channel x y - log x channel. If you give it a number for y, it will log the y most recent "
                   "messages.\n\n"
                   "$clear x - Clears the x most recent messages. If no x is given, clears the entire current channel. "
                   "Only those with permission can use this command.\n\n"
                   "$clear_channel x y - clears x channel. If you give it a number for y, it will clear the y most "
                   "recent messages. Only those with permission can use this command."
                   "```")


@client.event
async def on_close():
    await client.change_presence(status=discord.Status.offline, activity=None)


client.run('BOT_TOKEN')
